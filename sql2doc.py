import codecs
import re
import os
import natsort

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

TITLE = 'ShuoGG'
NORMAL_FONT = '微软雅黑'
NORMAL_FONT_SIZE = 10
TABLE_STYLE = 'Light Grid Accent 3'


class Field:
    def __init__(self, name='', type='', comment=''):
        self.name = name
        self.type = type
        self.comment = comment
        self.is_not_null = False
        self.is_key = False
        self.is_auto_increase = False

    def __repr__(self):
        return "Field(name=%s, type=%s, comment=%s)" % (self.name, self.type, self.comment)


class Table:
    def __init__(self):
        self.name = ''
        self.comment = ''
        self.fields = []

    def __repr__(self):
        return "Table(name=%s, comment=%s, field=%s)" % (self.name, self.comment, repr(self.fields))


def quick_mkdir(name):
    """
    当前目录下建一个文件夹
    :param name: 文件夹名称
    :return: 新建的文件夹的完整路径
    """
    new_directory = os.getcwd() + '\\' + name + "\\"
    if not os.path.exists(new_directory):
        try:
            os.mkdir(os.getcwd() + '\\' + name)
        except Exception as e:
            print(e)
    return new_directory


def get_suffixfiles_fullpath(suffix):
    """
    获取当前目录下所有.xxx文件的路径 (路径列表自然排序)
    :param suffix: 后缀如".sql" ".java"
    :return: list of str
    """
    sql_files = list(filter(lambda x: x.endswith(suffix), os.listdir(os.getcwd())))
    sqlFilesFullPath = list(map(lambda x: os.getcwd() + '\\' + x, sql_files))
    return natsort.natsorted(sqlFilesFullPath)


def get_file_content(file_path):
    """
    读取文件, 暂时只支持utf8和gbk编码的文件, 自动去除BOM
    :param file_path:
    :return: str
    """
    try:
        with open(file_path, encoding='utf-8') as f1:
            raw = f1.read()
            # 去掉BOM
            bom_head = raw.encode(encoding='utf-8')[:3]
            if bom_head == codecs.BOM_UTF8:
                raw = raw.encode(encoding='utf-8')[3:].decode(encoding='utf-8')
            return raw
    except Exception as e:
        with open(file_path, encoding='GBK') as f2:
            return f2.read()


def get_tables(sql_text):
    """
    sql文本转Table对象列表
    :param sql_text: sql文本
    :return: list of Table
    """
    tables = []
    ret = re.findall(r'(create\stable\s(.+)[^\(]+\(([^;]+)\);)', sql_text)
    for per_table_ret in ret:
        table_name = per_table_ret[1]
        table_body = per_table_ret[2]
        table_body_lines = list(map(lambda x: x.strip(), table_body.strip().splitlines()))
        new_table = Table()
        # 遍历( ... )里面的每一行
        for line in table_body_lines:
            key_names = []
            if 'primary key' in line:
                key_names.append(line[line.find('(') + 1:line.find(')')])
            if 'key ' not in line:
                ret_line = re.search(r'([\w_]+)\s+([^\s]+)\s?(.*)', line)
                field_name = ret_line.group(1)
                field_type = ret_line.group(2)
                field_tail = ret_line.group(3)
                field_comment = '--'
                if 'comment ' in field_tail:
                    field_comment = re.search(r'[\w\s]+\'(.+)\'', line).group(1)
                new_field = Field(field_name, field_type, field_comment)
                if 'not null' in line:
                    new_field.is_not_null = True
                if 'auto_increment' in line:
                    new_field.is_auto_increase = True
                new_table.fields.append(new_field)
        # 处理primary key
        for per_field in new_table.fields:
            if per_field.name in key_names:
                per_field.is_key = True
        table_comment = 'null'
        ret_comment = re.search(r'alter\stable\s%s.+\'(.+)\'' % table_name, sql_text)
        if ret_comment is not None:
            table_comment = ret_comment.group(1)
        new_table.comment = table_comment
        new_table.name = table_name
        tables.append(new_table)
    return tables


def gen_docx(tables, doc_name='sql.docx'):
    """
    通过list of Table生成docx
    :param tables: list of Table obj
    :param doc_name: file full path
    :return: bool
    """
    if tables is None or len(tables) == 0:
        print('Table list is None or Empty!')
        return False
    document = Document()
    # 样式设置
    document.styles['Normal'].font.name = NORMAL_FONT
    document.styles['Normal'].font.size = Pt(NORMAL_FONT_SIZE)
    # noinspection PyProtectedMember
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 1.Head生成
    head_runs = document.add_heading(TITLE, 0).runs
    for per_run in head_runs:
        per_run.font.name = NORMAL_FONT
        # noinspection PyProtectedMember
        per_run._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 2.Sql表格的生成
    for sql_table in tables:
        # 表Title生成
        tag_text = sql_table.name + '(' + sql_table.comment + ')' if not sql_table.comment == 'null' else sql_table.name
        label = document.add_paragraph()
        label_table_name = label.add_run(tag_text[0].upper() + tag_text[1:])
        label_table_name.font.size = Pt(14)
        # 表格生成
        row, col = len(sql_table.fields) + 1, 6
        docx_table = document.add_table(rows=row, cols=col, style=TABLE_STYLE)
        docx_table.cell(0, 0).text = '字段名称'
        docx_table.cell(0, 1).text = '字段注释'
        docx_table.cell(0, 2).text = '数据类型'
        docx_table.cell(0, 3).text = '关键字'
        docx_table.cell(0, 4).text = '不可空'
        docx_table.cell(0, 5).text = '自增'
        for field, row_index in zip(sql_table.fields, range(len(sql_table.fields))):
            docx_table.cell(row_index + 1, 0).text = field.name
            docx_table.cell(row_index + 1, 1).text = field.comment
            docx_table.cell(row_index + 1, 2).text = field.type
            docx_table.cell(row_index + 1, 3).text = str(field.is_key)
            docx_table.cell(row_index + 1, 4).text = str(field.is_not_null)
            docx_table.cell(row_index + 1, 5).text = str(field.is_auto_increase)
        document.add_paragraph("\n")
    document.save(doc_name)
    return True


def run_transfer():
    global TITLE
    path_list = get_suffixfiles_fullpath('.sql')
    dst_folder_path = quick_mkdir('1.new docx')
    if len(path_list) > 0:
        for sql_file_path in path_list:
            path, filename = os.path.split(sql_file_path)
            name, suffix = os.path.splitext(filename)
            sql_text = get_file_content(sql_file_path)
            table_list = get_tables(sql_text)
            docx_file_name = name + '.docx'
            TITLE = name
            if gen_docx(table_list, dst_folder_path + docx_file_name):
                print(docx_file_name + '转换成功!!')
            else:
                print(docx_file_name + '转换失败!!')


if __name__ == '__main__':
    run_transfer()
