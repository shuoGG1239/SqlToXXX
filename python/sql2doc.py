import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from util import sql_parser, file_util, system_util

TITLE = 'ShuoGG'
NORMAL_FONT = '微软雅黑'
NORMAL_FONT_SIZE = 10
TABLE_STYLE = 'Light Grid Accent 3'




def gen_docx(tables, doc_name='sql.docx'):
    """
    通过list of Table生成docx
    :param tables: list of Table obj
    :param doc_name: file full path
    :return: bool
    """
    if tables is None or len(tables) == 0:
        print('Table list is None or Empty!')
        return False
    document = Document()
    # 样式设置
    document.styles['Normal'].font.name = NORMAL_FONT
    document.styles['Normal'].font.size = Pt(NORMAL_FONT_SIZE)
    # noinspection PyProtectedMember
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 1.Head生成
    head_runs = document.add_heading(TITLE, 0).runs
    for per_run in head_runs:
        per_run.font.name = NORMAL_FONT
        # noinspection PyProtectedMember
        per_run._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 2.Sql表格的生成
    for sql_table in tables:
        # 表Title生成
        tag_text = sql_table.name + '(' + sql_table.comment + ')' if not sql_table.comment == 'null' else sql_table.name
        label = document.add_paragraph()
        label_table_name = label.add_run(tag_text[0].upper() + tag_text[1:])
        label_table_name.font.size = Pt(14)
        # 表格生成
        row, col = len(sql_table.fields) + 1, 6
        docx_table = document.add_table(rows=row, cols=col, style=TABLE_STYLE)
        docx_table.cell(0, 0).text = '字段名称'
        docx_table.cell(0, 1).text = '字段注释'
        docx_table.cell(0, 2).text = '数据类型'
        docx_table.cell(0, 3).text = '关键字'
        docx_table.cell(0, 4).text = '不可空'
        docx_table.cell(0, 5).text = '自增'
        for field, row_index in zip(sql_table.fields, range(len(sql_table.fields))):
            docx_table.cell(row_index + 1, 0).text = field.name
            docx_table.cell(row_index + 1, 1).text = field.comment
            docx_table.cell(row_index + 1, 2).text = field.type
            docx_table.cell(row_index + 1, 3).text = str(field.is_key)
            docx_table.cell(row_index + 1, 4).text = str(field.is_not_null)
            docx_table.cell(row_index + 1, 5).text = str(field.is_auto_increase)
        document.add_paragraph("\n")
    document.save(doc_name)
    return True


def run_transfer():
    global TITLE
    path_list = file_util.get_files_fullpath_curdir('.sql')
    dst_folder_path = file_util.quick_mkdir('1.new docx')
    if len(path_list) > 0:
        for sql_file_path in path_list:
            path, filename = os.path.split(sql_file_path)
            name, suffix = os.path.splitext(filename)
            sql_text = file_util.get_file_content(sql_file_path)
            table_list = sql_parser.get_tables(sql_text)
            docx_file_name = name + '.docx'
            TITLE = name
            if gen_docx(table_list, dst_folder_path + docx_file_name):
                print(docx_file_name + '转换成功!!')
            else:
                print(docx_file_name + '转换失败!!')


def sql_to_doc(text):
    global TITLE
    path_list = file_util.get_files_fullpath_curdir('.sql')
    dst_folder_path = file_util.quick_mkdir('1.new docx')
    name = 'new_new_new'
    table_list = sql_parser.get_tables(text)
    docx_file_name = name + '.docx'
    TITLE = name
    if gen_docx(table_list, dst_folder_path + docx_file_name):
        print(docx_file_name + '转换成功!!')
    else:
        print(docx_file_name + '转换失败!!')


if __name__ == '__main__':
    text = system_util.get_clipboard_text()
    if 'create table ' in text:
        # 直接对剪贴板的内容进行转换
        sql_to_doc(text)
    else:
        run_transfer()
